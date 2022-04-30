import docx


# Print DEBUGS
PRINT_TABLE = True

def getText(doc):
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def getTable(doc, tableIndex):
    # Data will be a list of rows represented as dictionaries
    # containing each row's data.
    table = doc.tables[tableIndex]
    data = []

    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if PRINT_TABLE:
            printText = list(text)
            print(f"{i}: {printText}")
            # for item in text:
            #     print(f"{item}")

        # Establish the mapping based on the first row
        # headers; these will become the keys of our dictionary
        if i == 0:
            keys = tuple(text)
            continue

        # Construct a dictionary for this row, mapping
        # keys to values for this row
        row_data = dict(zip(keys, text))
        data.append(row_data)

    return data


if __name__ == "__main__":
    doc_filename = f"case_reports/DRAFTCaseReport_template.docx"
    doc = docx.Document(doc_filename)

    print(f"{getText(doc)}")

    print(f"table output:")
    print(f"table 0:")
    print(f"{getTable(doc, 0)}")
    print(f"table 1:")
    print(f"{getTable(doc, 1)}")
    print(f"table 2:")
    print(f"{getTable(doc, 2)}")
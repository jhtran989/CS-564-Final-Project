import re

import docx
from my_table import *
from parse_doc import *


# Print DEBUGS
PRINT_TABLE = True
PRINT_BIOAFFINITY_TABLE = False

# FIXME: for now, working with only one report
# need to add JSON compatability
remainsId = 0
# method id/reference id should encompass ALL reports, but for demonstration,
# just use a baseline of 0 for each report
methodId = 0
referenceId = 0

def parseText(doc, tableTotal):
    fullText = getText(doc)

    # first line
    firstLine = re.search(r"RE:  Forensic anthropological analysis of (.*) "
                          r"discovered in ((.*), (.*)).", fullText)
    typeRemain = firstLine.group(1)
    country = firstLine.group(3)
    state = firstLine.group(4)

    print(f"First Line:")
    print(f"type of remain: {typeRemain}")
    print(f"country: {country}")
    print(f"state: {state}")
    print()

    # get sex
    sexLine = re.search(r"Biological Profile\n"
                          r"Sex: (.*)", fullText)
    sex = sexLine.group(1)

    print(f"Sex:")
    print(f"sex: {sex}")
    print()

    # get bioaffinity
    bioaffinityLine = re.search(r"Bioaffinity/Ancestry: (.*?)\n"
                                r"(.*?)\n\n", fullText, flags=re.DOTALL)
    bioaffinity = bioaffinityLine.group(1)
    bioaffinityNotes = bioaffinityLine.group(2)

    # add to table
    bioaffinityTable = tableTotal.getTable(TableEnum.BIOAFFINITY)
    bioaffinityTable.attributes[BioaffinityEnum.BIOAFFINITY_ESTIMATE] = \
        bioaffinity
    bioaffinityTable.attributes[BioaffinityEnum.BIOAFFINITY_NOTES] = \
        bioaffinityNotes

    print(f"Bioaffinity:")
    print(f"bioaffinity: {bioaffinity}")
    print(f"notes: {bioaffinityNotes}")
    print()

    # get age
    ageLine = re.search(r"Age: ((.*)-(.*)) years old", fullText)
    startAge = ageLine.group(2)
    endAge = ageLine.group(3)

    # add to table
    ageTable = tableTotal.getTable(TableEnum.AGE)
    ageTable.attributes[AgeEnum.AGE_START_ESTIMATE] = startAge
    ageTable.attributes[AgeEnum.AGE_START_ESTIMATE] = endAge

    print(f"Age:")
    print(f"start: {startAge}")
    print(f"end: {endAge}")
    print()

    # get stature
    statureLine = re.search(r"Stature: ((.*?)-(.*?)) inches\n"
                            r"(.*?)\n\n", fullText, flags=re.DOTALL)
    startStature = statureLine.group(2)
    endStature = statureLine.group(3)
    statureNotes = statureLine.group(4)

    # add to table
    statureTable = tableTotal.getTable(TableEnum.STATURE)
    statureTable.attributes[StatureEnum.STATURE_START_ESTIMATE] = startStature
    statureTable.attributes[StatureEnum.STATURE_END_ESTIMATE] = endStature
    statureTable.attributes[StatureEnum.STATURE_NOTES] = statureNotes

    print(f"Stature:")
    print(f"start: {startStature}")
    print(f"end: {endStature}")
    print(f"notes: {statureNotes}")
    print()

    # get individualizing characteristics
    individualizingCharLine = re.search(r"Individualizing characteristics:\n"
                                        r"(.*?)\n\n", fullText, flags=re.DOTALL)
    individualizingCharNotes = individualizingCharLine.group(1)

    # add to table
    icTable = tableTotal.getTable(TableEnum.INDIVIDUALIZING_CHARACTERISTICS)
    icTable.attributes[IndividualizingCharEnum.IC_NOTES] = \
        individualizingCharNotes

    print(f"Individualizing Characteristics:")
    print(f"notes: {individualizingCharNotes}")
    print()


def parseBiologicalProfileTable(doc, tableTotal):
    tableData = getTable(doc, 0)
    biologicalProfileLine = tableData[0][1]
    individualizingCharLine = tableData[1][1]
    perimortemTraumaLine = tableData[2][1]

    # biological profile values
    biologicalProfileValues = re.search(r"BIOLOGICAL PROFILE:\n"
                                            r"SEX:(.*)\n"
                                            r"AGE:(.*)\n"
                                            r"ANCESTRY:(.*)\n"
                                            r"STATURE:(.*)\n",
                                            biologicalProfileLine)

    sex = biologicalProfileValues.group(1).strip()
    age = biologicalProfileValues.group(2).strip()
    ancestry = biologicalProfileValues.group(3).strip()
    stature = biologicalProfileValues.group(4).strip()

    # add values to biological profile table
    bfTable = tableTotal.getTable(TableEnum.BIOLOGICAL_PROFILE)
    bfTable.attributes[BiologicalProfileEnum.SEX] = sex
    bfTable.attributes[BiologicalProfileEnum.AGE] = age
    bfTable.attributes[BiologicalProfileEnum.BIOAFFNITY] = ancestry
    bfTable.attributes[BiologicalProfileEnum.STATURE] = stature
    #tableTotal.addTable(bfTable)

    # individualizing characteristics values
    individualCharValues = re.search(r"INDIVIDUALIZING CHARACTERISTICS:\n"
                                     r"(.*)",
                                     individualizingCharLine, flags=re.DOTALL)

    individualCharNotes = individualCharValues.group(1).strip()

    # add values to biological profile table
    bfTable.attributes[BiologicalProfileEnum.INDIVIDUALIZING_CHARACTERISTICS] \
        = individualCharNotes

    # icTable = tableTotal.getTable(TableEnum.INDIVIDUALIZING_CHARACTERISTICS)
    # icTable.attributes[IndividualizingCharEnum.IC_NOTES] = individualCharNotes
    #tableTotal.addTable(icTable)

    # individualizing characteristics values
    perimortemTraumaValues = re.search(r"PERIMORTEM TRAUMA:\n"
                                       r"(.*)",
                                       perimortemTraumaLine, flags=re.DOTALL)

    perimortemTraumaNotes = perimortemTraumaValues.group(1).strip()

    # add values to perimortem trauma
    bfTable.attributes[BiologicalProfileEnum.PERIMORTEM_TRAUMA] \
        = perimortemTraumaNotes


# all parsing of table below are STATIC (tables remain at the same relative
# position in the doc -- i.e., no extra tables than listed in the template)
def parseBioaffinityTable(doc, tableIndex, tableTotal):
    global methodId
    global referenceId

    tableData = getTable(doc, tableIndex)
    bioaffinityMethodLines = tableData[1:]

    bioaffinityTable = tableTotal.getTable(TableEnum.BIOAFFINITY)
    bioaffinityId = bioaffinityTable.attributes[BioaffinityEnum.BIOAFFINITY_ID]

    referenceTableList = tableTotal.getTable(TableEnum.REFERENCE)

    for methodLine in bioaffinityMethodLines:
        print(f"{methodLine}")

        methodTable = Table(TableEnum.METHOD)
        methodTable.attributes[MethodEnum.PARENT_ID] = bioaffinityId
        methodTable.attributes[MethodEnum.METHOD_ID] = methodId
        methodId = methodId + 1

        # first attribute in the line
        methodName = methodLine[0]
        methodTable.attributes[MethodEnum.METHOD] = methodName

        # second attribute in the line
        comparisonGroup = methodLine[1]
        methodTable.attributes[MethodEnum.COMPARISON_GROUPS] = comparisonGroup

        # third attribute in the line
        bioaffinityEstimate = methodLine[2]
        methodTable.attributes[MethodEnum.ESTIMATE] = bioaffinityEstimate

        # fourth attribute in the line
        referenceName = methodLine[3]
        oldReferenceId = TableEncapsular.checkExistingReference(referenceName,
                                                             referenceTableList)

        if oldReferenceId is not None:
            methodTable.attributes[MethodEnum.REFERENCE_ID] = oldReferenceId
        else:
            referenceTable = Table(TableEnum.REFERENCE)
            referenceTable.attributes[ReferenceEnum.REFERENCE_ID] = referenceId
            referenceTable.attributes[ReferenceEnum.REFERENCE] = referenceName
            tableTotal.addTable(referenceTable)

            methodTable.attributes[MethodEnum.REFERENCE_ID] = referenceId
            referenceId = referenceId + 1

        tableTotal.addTable(methodTable)

    if PRINT_BIOAFFINITY_TABLE:
        print(f"{tableData}")


# all parsing of table below are STATIC (tables remain at the same relative
# position in the doc -- i.e., no extra tables than listed in the template)
def parseAgeTable(doc, tableIndex, tableTotal):
    global methodId
    global referenceId

    tableData = getTable(doc, tableIndex)
    bioaffinityMethodLines = tableData[1:]

    ageTable = tableTotal.getTable(TableEnum.AGE)
    ageId = ageTable.attributes[AgeEnum.AGE_ID]

    for methodLine in bioaffinityMethodLines:
        referenceTableList = tableTotal.getTable(TableEnum.REFERENCE)

        methodTable = Table(TableEnum.METHOD)
        methodTable.attributes[MethodEnum.PARENT_ID] = ageId
        methodTable.attributes[MethodEnum.METHOD_ID] = methodId
        methodId = methodId + 1

        # first attribute in the line
        methodName = methodLine[0]
        methodTable.attributes[MethodEnum.METHOD] = methodName

        # second attribute in the line
        comparisonGroup = methodLine[1]
        methodTable.attributes[MethodEnum.COMPARISON_GROUPS] = comparisonGroup

        # third attribute in the line
        bioaffinityEstimate = methodLine[2]
        methodTable.attributes[MethodEnum.ESTIMATE] = bioaffinityEstimate

        # fourth attribute in the line
        referenceName = methodLine[3]
        oldReferenceId = TableEncapsular.checkExistingReference(referenceName,
                                                             referenceTableList)

        if oldReferenceId is not None:
            methodTable.attributes[MethodEnum.REFERENCE_ID] = oldReferenceId
        else:
            referenceTable = Table(TableEnum.REFERENCE)
            referenceTable.attributes[ReferenceEnum.REFERENCE_ID] = referenceId
            referenceTable.attributes[ReferenceEnum.REFERENCE] = referenceName
            tableTotal.addTable(referenceTable)

            methodTable.attributes[MethodEnum.REFERENCE_ID] = referenceId
            referenceId = referenceId + 1

        tableTotal.addTable(methodTable)

    if PRINT_BIOAFFINITY_TABLE:
        print(f"{tableData}")


class TableEncapsular():
    def __init__(self):
        self.tables = []

    def addTable(self, table):
        self.tables.append(table)

    def getTable(self, tableEnum):
        """
        Should return a list of reference tables (since that should be the only
        one with MULTIPLE tuples)

        :param tableEnum:
        :return:
        """
        if tableEnum == TableEnum.REFERENCE:
            tableList = []
            for table in self.tables:
                if tableEnum == table.tableName:
                    tableList.append(table)

            return tableList
        else:
            for table in self.tables:
                if tableEnum == table.tableName:
                    return table

        return None

    @staticmethod
    def checkExistingReference(referenceName, referenceTableList):
        """
        Make sure to clean up methodName BEFORE doing this function

        :param referenceName:
        :param referenceTableList:
        :return:
        """
        for referenceTable in referenceTableList:
            #FIXME
            # print(f"{referenceName}")
            # print(f"table:"
            #       f" {referenceTable.attributes[ReferenceEnum.REFERENCE]}")

            if referenceName == \
                    referenceTable.attributes[ReferenceEnum.REFERENCE]:
                return referenceTable.attributes[ReferenceEnum.REFERENCE_ID]

        return None

    def __str__(self):
        totalString = ""
        for table in self.tables:
            totalString = totalString + f"{table.__str__()}\n"

        return totalString

if __name__ == "__main__":
    doc_filename = f"case_reports/DRAFTCaseReport_template_test.docx"
    doc = docx.Document(doc_filename)

    # initialize the tables
    tableTotal = TableEncapsular()
    for name, member in TableEnum.__members__.items():
        # print(f"{member}")
        # print(f"{TableEnum.REMAINS}")
        # a = TableEnum.REMAINS
        # print(f"{a}")
        # print(f"{member == TableEnum.METHOD}")
        if member != TableEnum.METHOD and member != TableEnum.REFERENCE:
            tableTotal.addTable(Table.tableFactory(member))

    #FIXME
    print(f"{tableTotal}")

    # set remains id
    remainsTable = tableTotal.getTable(TableEnum.REMAINS)
    remainsTable.attributes[Remains.REMAINS_ID] = remainsId

    #print(f"{getText(doc)}")
    parseText(doc, tableTotal)

    # Biological Profile table
    parseBiologicalProfileTable(doc, tableTotal)

    # Bioaffinity table
    parseBioaffinityTable(doc, 1, tableTotal)

    # Age table
    parseAgeTable(doc, 2, tableTotal)

    # print out final tuples...
    print()
    print(f"{tableTotal}")

    # print(f"table output:")
    # print(f"table 0:")
    # print(f"{getTable(doc, 0)}")
    # print(f"table 1:")
    # print(f"{getTable(doc, 1)}")
    # print(f"table 2:")
    # print(f"{getTable(doc, 2)}")